#!/usr/bin/env python3
"""
to_docx.py  --  Convert a Markdown report to a Word (.docx) document.

The analysis report is saved as both Markdown and Word. This makes the .docx from
the .md. Uses pandoc when available (best fidelity: headings, tables, bold/italic,
lists); otherwise falls back to a built-in python-docx renderer so a Word file is
ALWAYS produced (headings, paragraphs, bullet lists, and pipe tables).

Both pandoc and python-docx are installed by setup.sh into the conda env, so the
pandoc path is the normal one; the fallback is a safety net.

Usage:
  python3 to_docx.py --in AI_Analysis_Report.md --out AI_Analysis_Report.docx
"""
import sys, os, re, shutil, subprocess, argparse


def via_pandoc(md_path, out_path):
    pandoc = shutil.which("pandoc")
    if not pandoc:
        return False
    try:
        subprocess.run([pandoc, md_path, "-o", out_path, "--from", "gfm",
                        "--standalone"], check=True, capture_output=True, text=True)
        return os.path.exists(out_path)
    except subprocess.CalledProcessError as e:
        sys.stderr.write(f"[to_docx] pandoc failed ({e.stderr.strip()[:200]}); using fallback\n")
        return False


def _add_inline(paragraph, text):
    """Render minimal inline markdown (**bold**, *italic*, `code`) into runs."""
    for chunk in re.split(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            paragraph.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            paragraph.add_run(chunk[1:-1]).italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            r = paragraph.add_run(chunk[1:-1]); r.font.name = "Courier New"
        else:
            paragraph.add_run(chunk)


def via_python_docx(md_path, out_path):
    try:
        from docx import Document
    except ImportError:
        return False
    doc = Document()
    lines = open(md_path, encoding="utf-8").read().splitlines()
    i, n = 0, len(lines)
    while i < n:
        line = lines[i].rstrip()
        # pipe table: header row, separator row, then body rows
        if line.startswith("|") and i + 1 < n and re.match(r"^\s*\|[ \-:|]+\|\s*$", lines[i + 1]):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells); i += 1
            rows = [r for j, r in enumerate(rows) if j != 1]  # drop the --- separator
            if rows:
                t = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                t.style = "Light Grid Accent 1"
                for ri, r in enumerate(rows):
                    for ci, c in enumerate(r):
                        t.rows[ri].cells[ci].text = re.sub(r"[*`]", "", c)
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            doc.add_heading(m.group(2).strip(), level=min(len(m.group(1)), 4))
        elif re.match(r"^\s*[-*+]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, re.sub(r"^\s*[-*+]\s+", "", line))
        elif re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, re.sub(r"^\s*\d+\.\s+", "", line))
        elif line.strip() in ("---", "___", "***"):
            pass  # horizontal rule -> skip
        elif line.strip():
            _add_inline(doc.add_paragraph(), line)
        i += 1
    doc.save(out_path)
    return os.path.exists(out_path)


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--in", dest="inp", required=True)
    ap.add_argument("--out", required=True)
    a = ap.parse_args()
    if not os.path.exists(a.inp):
        sys.exit(f"input not found: {a.inp}")

    if via_pandoc(a.inp, a.out):
        print(f"[to_docx] wrote {a.out} (pandoc)")
    elif via_python_docx(a.inp, a.out):
        print(f"[to_docx] wrote {a.out} (python-docx fallback)")
    else:
        sys.exit("Could not convert to .docx: neither pandoc nor python-docx is available. "
                 "Re-run setup.sh to install them.")


if __name__ == "__main__":
    main()
